from flask import Flask, request, render_template, send_file
from docx import Document
from docx.shared import Inches
import os

app = Flask(__name__)


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/generate", methods=["POST"])
def generate_docx():
    # Collect form data
    data = {
        "output": request.form["output"],
        "vulnerability_title": request.form["title"],
        "severity": request.form["severity"],
        "status": request.form["status"],
        "cvss_score": request.form["cvss_score"],
        "cwe_id": request.form["cwe_id"],
        "owasp_category": request.form["owasp_category"],
        "description": request.form["description"],
        "impact": request.form["impact"],
        "affected_url": request.form["affected_url"],
        "reference_url": request.form["reference_url"],
        "recommendations": request.form["recommendations"],
        "steps": request.form.getlist("steps[]"),
    }

    # Handle file uploads
    pictures = request.files.getlist("pictures[]")
    picture_paths = []
    for i, picture in enumerate(pictures):
        if picture.filename != "":
            picture_path = os.path.join(
                "uploads", f"{data['vulnerability_title']}_poc_{i+1}.jpg"
            )
            picture.save(picture_path)
            picture_paths.append(picture_path)
        else:
            picture_paths.append(None)

    # Load the template and populate it
    template_path = "./template.docx"
    doc = Document(template_path)

    # making steps object with images
    steps_procedure = {}

    for i, step in enumerate(data["steps"]):
        steps_procedure[step] = picture_paths[i]

    # Example for populating fields
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    if key.lower() in cell.text and type(value) == str:
                        cell.text = cell.text.replace(key.lower(), value)
                if "steps_to_reproduce" in cell.text:
                    cell.text = cell.text.replace("steps_to_reproduce", "")
                    for i, step in enumerate(data["steps"]):
                        p = cell.add_paragraph(f"Step {i+1}: {step}")
                        if picture_paths[i]:
                            cell.add_paragraph()  # Add an empty paragraph for spacing
                            run = cell.add_paragraph().add_run()
                            run.add_picture(
                                picture_paths[i], width=Inches(6.5)
                            )  # Set the width to fit the page

    # Save the updated document
    output_path = data["output"] + ".docx"
    doc.save(output_path)

    return send_file(output_path, as_attachment=True)


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=80, debug=True)
